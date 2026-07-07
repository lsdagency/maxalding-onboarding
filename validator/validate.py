"""
File scanners and orchestration for the Maxalding QA gate.

Turns each input file into a list of Segment objects, then runs the deterministic
checks over them. Supports:

  .py            build-script string literals (character checks only)
  .md / .txt     documentation, treated as rule statements (character checks only)
  .docx          generated deliverables (full checks, header no-date check)
  .xlsx          the Creative Plan (full checks, tab-aware length and hook checks)

python-docx and openpyxl are imported lazily so the character and text checks
still work in environments where they are not installed.
"""

from __future__ import annotations

import ast
import os

from . import rules
from .checks import (
    Segment,
    Result,
    run_segment_checks,
    check_naming,
    is_rule_statement,
)

TEXT_EXTS = {".md", ".txt", ".markdown"}
CODE_EXTS = {".py"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx"}
SCANNABLE = TEXT_EXTS | CODE_EXTS | DOCX_EXTS | XLSX_EXTS

# Directories never scanned by default (they hold the rule definitions
# themselves and would produce documentation-only matches).
DEFAULT_SKIP_DIRS = {".git", "__pycache__", ".venv", "venv", "node_modules"}


# ---------------------------------------------------------------------------
# Scanners
# ---------------------------------------------------------------------------
def scan_py(path: str) -> list:
    """Every string literal in a build script (character checks only)."""
    with open(path, "r", encoding="utf-8") as fh:
        source = fh.read()
    segments = []
    try:
        tree = ast.parse(source)
    except SyntaxError:
        # Fall back to whole-file text so a broken script still gets char checks.
        return [Segment(source, f"{path}", kind="code_string")]
    for node in ast.walk(tree):
        if isinstance(node, ast.Constant) and isinstance(node.value, str):
            line = getattr(node, "lineno", "?")
            segments.append(
                Segment(node.value, f"{path}:{line}", kind="code_string")
            )
    return segments


def scan_text(path: str) -> list:
    """Markdown and text docs: char checks only (they quote banned items)."""
    segments = []
    with open(path, "r", encoding="utf-8") as fh:
        for i, line in enumerate(fh, start=1):
            stripped = line.rstrip("\n")
            if stripped.strip():
                segments.append(
                    Segment(stripped, f"{path}:{i}", kind="rule_statement")
                )
    return segments


def scan_docx(path: str) -> list:
    from docx import Document  # lazy import

    doc = Document(path)
    segments = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        loc = f"{path} [p{idx}]"
        if idx < rules.HEADER_SCAN_PARAGRAPHS:
            segments.append(Segment(text, loc, kind="header"))
        kind = "rule_statement" if is_rule_statement(text) else "content"
        segments.append(Segment(text, loc, kind=kind))
    # Tables (CRM, landing page grids) are content too.
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                if not text.strip():
                    continue
                loc = f"{path} [table{t_idx} r{r_idx} c{c_idx}]"
                kind = "rule_statement" if is_rule_statement(text) else "content"
                segments.append(Segment(text, loc, kind=kind))
    return segments


def _col_letter(idx: int) -> str:
    from openpyxl.utils import get_column_letter

    return get_column_letter(idx)


def scan_xlsx(path: str) -> list:
    import openpyxl  # lazy import

    wb = openpyxl.load_workbook(path, data_only=False)
    segments = []
    for ws in wb.worksheets:
        sheet = ws.title
        sheet_lower = sheet.lower()

        # Build an override map of (row, col) -> kind for tab-aware tagging.
        overrides = {}

        # Locate a header row by scanning the first rows for known labels.
        header_row = None
        headers = {}  # column index -> label (upper)
        for r in range(1, min(ws.max_row, 6) + 1):
            row_labels = {}
            for c in range(1, ws.max_column + 1):
                v = ws.cell(row=r, column=c).value
                if isinstance(v, str):
                    row_labels[c] = v.strip().upper()
            if "HOOK" in row_labels.values() or "POST COPY" in row_labels.values():
                header_row = r
                headers = row_labels
                break

        if header_row:
            label_to_col = {lbl: col for col, lbl in headers.items()}

            # Creative Tracker: tag HOOK cells as static_hook / video_hook by
            # the FORMAT column on the same row.
            if "HOOK" in label_to_col and "FORMAT" in label_to_col:
                hook_col = label_to_col["HOOK"]
                fmt_col = label_to_col["FORMAT"]
                for r in range(header_row + 1, ws.max_row + 1):
                    fmt = ws.cell(row=r, column=fmt_col).value
                    if not isinstance(fmt, str):
                        continue
                    fu = fmt.strip().upper()
                    if fu not in ("STATIC", "VIDEO"):
                        continue  # existing-post or other rows are not hook-checked
                    kind = "static_hook" if fu == "STATIC" else "video_hook"
                    overrides[(r, hook_col)] = ("hooklines", kind)

            # Creative Tracker: the DATE column must be blank. Tag any non-empty
            # DATE cell so the QA gate flags a stamped date (no implied deadlines).
            if "DATE" in label_to_col:
                date_col = label_to_col["DATE"]
                for r in range(header_row + 1, ws.max_row + 1):
                    overrides[(r, date_col)] = ("kind", "tracker_date")

            # AD COPY: tag the POST COPY column.
            if "POST COPY" in label_to_col:
                pc_col = label_to_col["POST COPY"]
                for r in range(header_row + 1, ws.max_row + 1):
                    v = ws.cell(row=r, column=pc_col).value
                    if isinstance(v, str) and v.strip():
                        overrides[(r, pc_col)] = ("kind", "post_copy")

        # AD COPY headline / description blocks: find label cells anywhere.
        for r in range(1, ws.max_row + 1):
            for c in range(1, ws.max_column + 1):
                v = ws.cell(row=r, column=c).value
                if not isinstance(v, str):
                    continue
                label = v.strip().lower()
                target = None
                if label in ("headline", "headlines"):
                    target = "headline"
                elif label in ("description", "descriptions"):
                    target = "description"
                if target:
                    # Tag subsequent non-empty cells in the same column.
                    rr = r + 1
                    while rr <= ws.max_row:
                        nv = ws.cell(row=rr, column=c).value
                        if not (isinstance(nv, str) and nv.strip()):
                            break
                        overrides[(rr, c)] = ("kind", target)
                        rr += 1

        # Emit segments for every non-empty string cell, applying overrides.
        for r in range(1, ws.max_row + 1):
            for c in range(1, ws.max_column + 1):
                v = ws.cell(row=r, column=c).value
                if not isinstance(v, str) or not v.strip():
                    continue
                coord = f"{_col_letter(c)}{r}"
                loc = f"{path} [{sheet}!{coord}]"
                ov = overrides.get((r, c))
                if ov and ov[0] == "hooklines":
                    # One segment per hook line for word-count checking, plus a
                    # whole-cell content segment for char/meta/premium checks.
                    for line in v.splitlines():
                        if line.strip():
                            segments.append(Segment(line.strip(), loc, kind=ov[1]))
                    segments.append(Segment(v, loc, kind="content"))
                elif ov and ov[0] == "kind":
                    segments.append(Segment(v, loc, kind=ov[1]))
                else:
                    kind = "rule_statement" if is_rule_statement(v) else "content"
                    segments.append(Segment(v, loc, kind=kind))
    return segments


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------
def _dedupe(violations: list) -> list:
    seen = set()
    out = []
    for v in violations:
        key = (v.rule, v.location, v.message, v.snippet)
        if key not in seen:
            seen.add(key)
            out.append(v)
    return out


def scan_file(path: str, check_names: bool = True) -> list:
    ext = os.path.splitext(path)[1].lower()
    violations = []
    if ext in CODE_EXTS:
        segments = scan_py(path)
    elif ext in TEXT_EXTS:
        segments = scan_text(path)
    elif ext in DOCX_EXTS:
        segments = scan_docx(path)
        if check_names:
            violations.extend(check_naming(os.path.basename(path)))
    elif ext in XLSX_EXTS:
        segments = scan_xlsx(path)
        if check_names:
            violations.extend(check_naming(os.path.basename(path)))
    else:
        return []
    violations.extend(run_segment_checks(segments))
    return violations


def _iter_files(paths, skip_dirs):
    for p in paths:
        if os.path.isdir(p):
            for root, dirs, files in os.walk(p):
                dirs[:] = [d for d in dirs if d not in skip_dirs]
                for name in files:
                    ext = os.path.splitext(name)[1].lower()
                    if ext in SCANNABLE and not name.startswith("~$"):
                        yield os.path.join(root, name)
        elif os.path.isfile(p):
            yield p


def validate_paths(paths, check_names: bool = True, skip_dirs=None, skip_rules=None) -> Result:
    """Scan paths and return a Result.

    skip_rules: an optional iterable of rule ids to waive (their violations are
    dropped from the result). Used for per-client exceptions, for example a
    low-ticket free-trial client that legitimately names a "free" offer, which
    waives "premium-framing". See build.build_all and the premium_lead_magnet
    client flag.
    """
    skip_dirs = DEFAULT_SKIP_DIRS if skip_dirs is None else skip_dirs
    skip_rules = set(skip_rules or ())
    result = Result()
    for path in _iter_files(paths, skip_dirs):
        try:
            result.violations.extend(scan_file(path, check_names=check_names))
        except Exception as exc:  # keep scanning the rest of the set
            result.violations.append(_scan_error(path, exc))
    result.violations = _dedupe(result.violations)
    if skip_rules:
        result.violations = [v for v in result.violations if v.rule not in skip_rules]
    return result


def _scan_error(path, exc):
    from .checks import Violation

    return Violation(
        rule="scan-error",
        message=f"Could not scan file: {exc}",
        location=path,
        severity="warn",
    )


def validate_text(text: str, kind: str = "content", location: str = "<text>") -> Result:
    """Convenience entry point for unit tests and the eval harness."""
    result = Result()
    result.violations = _dedupe(run_segment_checks([Segment(text, location, kind=kind)]))
    return result
