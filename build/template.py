"""
Shared template module for every Maxalding deliverable.

One place defines the font, the document header, the logo insertion and all
styling so a brand change is a single edit. The Creative Plan XLSX generator and
the four .docx generators all build on these helpers.

Writable, per-agency state (the logo, the feedback log, brand kits) lives OUTSIDE
this read-only package, in a workspace folder. Resolve it with resolve_workspace()
and never write into the package at runtime.

House rules enforced here:
  - Helvetica Neue on every run (python-docx falls back to Cambria otherwise,
    Feedback round 8), set explicitly regardless of size or style.
  - Document header carries CLIENT, CAMPAIGN, PREPARED BY only. No date.
  - Main heading 16 bold, subheadings 14 bold, body and script lines 10.
  - Spoken script lines in grey #787878, direction notes italic grey #505050.
  - XLSX: black header rows (white bold size 8), body size 8 #434343, wrap, top.
"""

from __future__ import annotations

import os

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------
FONT = "Helvetica Neue"

# docx sizes (points)
SIZE_MAIN_HEADING = 16
SIZE_SUBHEADING = 14
SIZE_BODY = 10

# docx colours (R, G, B)
COLOR_BLACK = (0, 0, 0)
COLOR_SCRIPT_GREY = (120, 120, 120)     # #787878 spoken lines
COLOR_DIRECTION_GREY = (80, 80, 80)     # #505050 direction notes

# xlsx styling
XLSX_FONT = "Helvetica Neue"
XLSX_HEADER_BG = "000000"
XLSX_HEADER_FG = "FFFFFF"
XLSX_BODY_FG = "434343"
XLSX_FONT_SIZE = 8
XLSX_MIN_ROW_HEIGHT = 45
XLSX_TRACKER_ROW_HEIGHT = 135  # within the 120 to 150pt band

LOGO_FILENAME = "Maxalding Logo.png"
DEFAULT_LOGO_WIDTH_CM = 1.0


# ---------------------------------------------------------------------------
# Workspace resolution (writable state lives outside the package)
# ---------------------------------------------------------------------------
def resolve_workspace(workspace: str | None = None) -> str:
    """
    Resolve the shared agency workspace folder.

    Order of precedence:
      1. explicit `workspace` argument
      2. MAXALDING_WORKSPACE environment variable
      3. ./maxalding-workspace under the current working directory
    """
    if workspace:
        return os.path.abspath(workspace)
    env = os.environ.get("MAXALDING_WORKSPACE")
    if env:
        return os.path.abspath(env)
    return os.path.abspath(os.path.join(os.getcwd(), "maxalding-workspace"))


def logo_path(workspace: str | None = None) -> str | None:
    ws = resolve_workspace(workspace)
    candidate = os.path.join(ws, LOGO_FILENAME)
    return candidate if os.path.isfile(candidate) else None


def deliverable_filename(client_business_name: str, deliverable: str) -> str:
    """MAXALDING - [Client Business Name] - [Deliverable].[ext]"""
    ext = {
        "Creative Plan": "xlsx",
        "Video Ad Scripts": "docx",
        "VSL Script": "docx",
        "Landing Page Copy": "docx",
        "CRM Automation": "docx",
    }[deliverable]
    return f"MAXALDING - {client_business_name} - {deliverable}.{ext}"


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------
def _set_run_font(run, size=None, bold=False, italic=False, color=None):
    """Force Helvetica Neue and styling on a single run. Always call this."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    run.font.name = FONT
    # Belt and braces: set the east-asian and complex-script names too, so the
    # theme default (Cambria) can never leak in.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def new_document():
    """A document whose default style is already Helvetica Neue."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(SIZE_BODY)
    return doc


def add_logo(doc, workspace=None, width_cm=DEFAULT_LOGO_WIDTH_CM) -> bool:
    """
    Insert the shared logo at the top of the document from the workspace folder.
    Returns True if inserted, False if the logo file was not found.
    """
    from docx.shared import Cm

    path = logo_path(workspace)
    if not path:
        return False
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return True


def add_main_heading(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, size=SIZE_MAIN_HEADING, bold=True, color=COLOR_BLACK)
    return para


def add_subheading(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, size=SIZE_SUBHEADING, bold=True, color=COLOR_BLACK)
    return para


def add_label_line(doc, label, value):
    """Bold label, unbolded value, size 10. Used for CLIENT / CAMPAIGN etc."""
    para = doc.add_paragraph()
    lab = para.add_run(f"{label}: ")
    _set_run_font(lab, size=SIZE_BODY, bold=True, color=COLOR_BLACK)
    val = para.add_run(value)
    _set_run_font(val, size=SIZE_BODY, bold=False, color=COLOR_BLACK)
    return para


def add_header_block(doc, client, campaign, prepared_by, workspace=None,
                     logo_width_cm=DEFAULT_LOGO_WIDTH_CM):
    """
    Standard document header for all four decks. Logo, main heading, then the
    three header labels. No date, ever (Feedback 2026-06-18).
    """
    add_logo(doc, workspace=workspace, width_cm=logo_width_cm)
    add_main_heading(doc, "MAXALDING AGENCY")
    add_label_line(doc, "CLIENT", client)
    add_label_line(doc, "CAMPAIGN", campaign)
    add_label_line(doc, "PREPARED BY", prepared_by)


def add_body(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, size=SIZE_BODY, bold=bold, color=COLOR_BLACK)
    return para


def add_direction(doc, text):
    """Italic direction note in dark grey."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, size=SIZE_BODY, italic=True, color=COLOR_DIRECTION_GREY)
    return para


def add_script_line(doc, text):
    """Spoken line rendered in light grey, wrapped in quotation marks."""
    para = doc.add_paragraph()
    quoted = text if text.startswith('"') else f'"{text}"'
    run = para.add_run(quoted)
    _set_run_font(run, size=SIZE_BODY, color=COLOR_SCRIPT_GREY)
    return para


def add_timestamp(doc, label):
    """Timestamp / section label, then a blank line for breathing room."""
    para = doc.add_paragraph()
    run = para.add_run(label)
    _set_run_font(run, size=SIZE_BODY, bold=True, color=COLOR_BLACK)
    doc.add_paragraph()  # blank line after every timestamp (Feedback round 7)
    return para


def add_divider(doc):
    para = doc.add_paragraph()
    run = para.add_run("-" * 40)
    _set_run_font(run, size=SIZE_BODY, color=COLOR_DIRECTION_GREY)
    return para


# ---------------------------------------------------------------------------
# xlsx helpers
# ---------------------------------------------------------------------------
def xlsx_font(bold=False, header=False):
    from openpyxl.styles import Font

    color = XLSX_HEADER_FG if header else XLSX_BODY_FG
    return Font(name=XLSX_FONT, size=XLSX_FONT_SIZE, bold=bold, color=color)


def xlsx_header_fill():
    from openpyxl.styles import PatternFill

    return PatternFill(start_color=XLSX_HEADER_BG, end_color=XLSX_HEADER_BG,
                       fill_type="solid")


def style_header_cell(cell):
    from openpyxl.styles import Alignment

    cell.font = xlsx_font(bold=True, header=True)
    cell.fill = xlsx_header_fill()
    cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")


def style_body_cell(cell):
    from openpyxl.styles import Alignment

    cell.font = xlsx_font(bold=False, header=False)
    cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")


def set_column_widths(ws, widths: dict):
    """widths: {column_letter: width}"""
    for letter, width in widths.items():
        ws.column_dimensions[letter].width = width


def add_dropdown(ws, cell_range, options: list, default=None):
    """Attach a list data-validation dropdown to a cell range."""
    from openpyxl.worksheet.datavalidation import DataValidation

    formula = '"' + ",".join(options) + '"'
    dv = DataValidation(type="list", formula1=formula, allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(cell_range)
    return dv
